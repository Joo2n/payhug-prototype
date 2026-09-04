#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""기호 정리표 — 워드와 HTML 두 벌.

원고는 아티팩트 HTML(ceo_review.html) 한 곳이다. 그 파일을 파싱해 같은 차례·같은
표로 워드를 짜고, HTML 은 머리를 씌워 그대로 낸다. 두 산출물이 같은 말을 한다.

산출 ~/Downloads/payhug_용어정의서/기호정리표_YYYYMMDD_HHMM.docx
     ~/Downloads/payhug_용어정의서/기호정리표_YYYYMMDD_HHMM.html
"""
import html as _html
import io
import os
import re
import sys
from datetime import datetime
from html.parser import HTMLParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PIPE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PIPE)
import subscript

SRC = os.environ.get("SYMREVIEW_SRC") or os.path.join(
    "/private/tmp/claude-501/-Users-semi-cursor-payhug",
    "9aed3429-fc00-4785-9abd-c254e437cf03/scratchpad/ceo_review.html")
OUT = os.path.expanduser("~/Downloads/payhug_용어정의서")
HAN, MONO = "맑은 고딕", "D2Coding"
INK = RGBColor(0x1C, 0x1B, 0x19)
MUTE = RGBColor(0x6E, 0x6A, 0x63)
ACC = RGBColor(0x2C, 0x44, 0x70)
WARN = RGBColor(0xA6, 0x3D, 0x2F)
OK = RGBColor(0x2F, 0x6B, 0x4F)
HEADFILL, GRPFILL, BOXFILL = "F4F1EC", "EFEBE4", "F7F5F1"


# ── 아티팩트 파싱 ──────────────────────────────────────────────
class Reader(HTMLParser):
    """블록을 차례대로 뽑는다. 표 칸은 (글자, 색갈래) 로 담는다."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self.stack = []
        self.buf = []
        self.tone = None
        self.tbl = None
        self.row = None
        self.rowkind = "body"
        self.cellspan = 1
        self.spans = []

    # 글자를 모으는 자리인가
    def _in(self, *names):
        return any(t[0] in names for t in self.stack)

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        cls = a.get("class", "")
        self.stack.append((tag, cls))
        if tag == "table":
            self.tbl = {"head": [], "rows": []}
        elif tag == "tr":
            self.row = []
            k = cls.split()
            self.rowkind = "__head" if "__head" in k else ("grp" if "grp" in k else "body")
        elif tag in ("td", "th"):
            self.buf, self.tone = [], None
            self.cellspan = int(a.get("colspan", 1))
            if "n" in cls.split() or "f" in cls.split():
                self.tone = "mono"
        elif tag in ("h1", "h2", "h3", "p", "div"):
            if tag == "div" and "formula" not in cls:
                return
            self.buf = []
        elif tag == "sub":
            self.buf.append("\x01")            # 아래첨자 여는 표시
        elif tag == "br":
            self.buf.append("\n")
        elif tag == "span":
            k = cls.split()
            mark = ("\x02" if "new" in k else "\x03" if "old" in k
                    else "\x04" if ("ok" in k or "p-ok" in k)
                    else "\x05" if ("hi" in k or "p-no" in k) else "")
            self.spans.append(bool(mark))
            if mark:
                self.buf.append(mark)

    def handle_data(self, d):
        if self.stack:
            self.buf.append(d)

    def handle_endtag(self, tag):
        cls = ""
        for i in range(len(self.stack) - 1, -1, -1):
            if self.stack[i][0] == tag:
                cls = self.stack[i][1]
                del self.stack[i:]
                break
        if tag == "sub":
            self.buf.append("\x81")            # 아래첨자 닫는 표시
            return
        if tag == "span":
            if self.spans and self.spans.pop():
                self.buf.append("\x80")
            return
        txt = self._text()
        if tag in ("td", "th"):
            self.row.append((txt, self.tone, self.cellspan))
        elif tag == "tr":
            if self.tbl is not None:
                self.tbl["rows"].append((self.rowkind, self.row))
            self.row = None
        elif tag == "table":
            self.blocks.append(("table", self.tbl))
            self.tbl = None
        elif tag in ("h1", "h2", "h3"):
            self.blocks.append((tag, txt))
        elif tag == "p":
            self.blocks.append(("note" if "note" in cls.split() else "p", txt))
        elif tag == "div" and "formula" in cls.split():
            self.blocks.append(("formula", txt))

    def handle_startendtag(self, tag, attrs):
        if tag == "br":
            self.buf.append("\n")

    def _text(self):
        t = "".join(self.buf)
        self.buf = []
        return t


def parse(path):
    src = io.open(path, encoding="utf-8").read()
    body = src[src.index('<div class="wrap">'):]
    # thead 안의 tr 을 머리행으로 표시해 둔다
    body = re.sub(r"<thead>(.*?)</thead>",
                  lambda m: m.group(1).replace("<tr>", '<tr class="__head">'),
                  body, flags=re.S)
    r = Reader()
    r.feed(body)
    # 머리행 갈래를 되살린다
    out = []
    for kind, val in r.blocks:
        if kind == "table":
            head, rows = [], []
            for rk, cells in val["rows"]:
                if rk == "__head":
                    head = cells
                else:
                    rows.append((rk, cells))
            out.append(("table", {"head": head, "rows": rows}))
        else:
            out.append((kind, val))
    return out


# ── 워드 조판 ─────────────────────────────────────────────────
def font(run, name=HAN, size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for ax in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rf.set(qn(ax), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


TONE = {"\x02": (ACC, True), "\x03": (WARN, False),
        "\x04": (OK, True), "\x05": (WARN, True)}


def write(p, t, name=HAN, size=None, bold=None, color=INK):
    """표시 문자를 풀어 run 으로 넣는다. \x01…\x81 은 아래첨자."""
    sub = False
    cur = (color, bold)
    stack = []
    chunk = ""

    def flush():
        nonlocal chunk
        if not chunk:
            return
        r = font(p.add_run(chunk), name, size, cur[1], cur[0])
        if sub:
            r.font.subscript = True
        chunk = ""

    for ch in t:
        if ch == "\x01":
            flush(); sub = True
        elif ch == "\x81":
            flush(); sub = False
        elif ch in TONE:
            flush(); stack.append(cur); cur = TONE[ch]
        elif ch == "\x80":
            flush(); cur = stack.pop() if stack else (color, bold)
        elif ch == "\n":
            flush(); p.add_run().add_break()
        else:
            chunk += ch
    flush()
    return p


def plain(t):
    return re.sub(r"[\x01-\x05\x80\x81]", "", t)


def para(doc, t="", size=10.5, bold=False, color=INK, name=HAN,
         before=0, after=4, line=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), line
    if t:
        write(p, t, name, size, bold, color)
    return p


def shade(el, hx):
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:fill"), hx)
    el.append(s)


def box(doc, t, size=9, name=MONO):
    """산식 블록 — 등폭에 옅은 바탕."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(4), Pt(8), 1.45
    pf.left_indent = Cm(0.3)
    shade(p._p.get_or_add_pPr(), BOXFILL)
    write(p, t.strip("\n"), name, size)
    return p


def cell(c, t, size=9, bold=False, color=INK, name=HAN, fill=None):
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    if fill:
        shade(c._tc.get_or_add_tcPr(), fill)
    write(p, t, name, size, bold, color)


def widths_for(head, rows):
    """열 폭 — 글자 수에 맞춰 16.5cm 를 나눈다."""
    n = len(head)
    w = [0] * n
    for _rk, cells in rows:
        if len(cells) != n:
            continue
        for i, (t, _tone, _sp) in enumerate(cells):
            w[i] = max(w[i], min(len(plain(t)), 60))
    for i, (t, _tone, _sp) in enumerate(head):
        w[i] = max(w[i], len(plain(t)))
    tot = sum(w) or 1
    return [max(1.7, 16.5 * x / tot) for x in w]


def put_table(doc, tb):
    head, rows = tb["head"], tb["rows"]
    n = len(head)
    if not n:
        return
    w = widths_for(head, rows)
    t = doc.add_table(rows=1, cols=n)
    t.style = "Table Grid"
    for i, (h, _tone, _sp) in enumerate(head):
        t.rows[0].cells[i].width = Cm(w[i])
        cell(t.rows[0].cells[i], h, 8.5, True, MUTE, fill=HEADFILL)
    for rk, cells in rows:
        rr = t.add_row()
        if len(cells) == 1 and cells[0][2] > 1:           # 그룹 행
            merged = rr.cells[0]
            for j in range(1, n):
                merged = merged.merge(rr.cells[j])
            cell(merged, cells[0][0], 8.5, True, MUTE, fill=GRPFILL)
            continue
        for i, (v, tone, _sp) in enumerate(cells[:n]):
            rr.cells[i].width = Cm(w[i])
            cell(rr.cells[i], v, 9, name=(MONO if tone == "mono" else HAN))
    para(doc, "", after=6)


def docx(blocks, path):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)
    nm = doc.styles["Normal"]
    nm.font.size = Pt(10.5)
    rpr = nm.element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rpr.append(rf)
    for ax in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rf.set(qn(ax), HAN)

    for kind, val in blocks:
        if kind == "h1":
            para(doc, val, 20, True, after=14)
        elif kind == "h2":
            para(doc, val, 14, True, before=16, after=6)
        elif kind == "h3":
            para(doc, val, 11, True, before=10, after=4)
        elif kind == "p":
            para(doc, val, 10, after=6)
        elif kind == "note":
            para(doc, val, 9, color=MUTE, after=8)
        elif kind == "formula":
            box(doc, val)
        elif kind == "table":
            put_table(doc, val)
    doc.save(path)
    return path


# ── HTML ──────────────────────────────────────────────────────
HEAD = """<!doctype html>
<html lang="ko"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>기호 정리표</title>
"""


def html_doc(path):
    src = io.open(path, encoding="utf-8").read()
    i = src.index("<style>")
    j = src.index("<div class=\"wrap\">")
    link = re.search(r'<link[^>]+fonts\.googleapis[^>]+>', src)
    return (HEAD + (link.group(0) + "\n" if link else "")
            + src[i:j] + "</head><body>\n" + src[j:] + "\n</body></html>\n")


def main():
    if not os.path.exists(SRC):
        sys.exit("원고 없음 — %s" % SRC)
    os.makedirs(OUT, exist_ok=True)
    stamp = (os.environ.get("SYMREVIEW_STAMP")
             or datetime.now().strftime("%Y%m%d_%H%M"))
    blocks = parse(SRC)
    dp = docx(blocks, os.path.join(OUT, "기호정리표_%s.docx" % stamp))
    hp = os.path.join(OUT, "기호정리표_%s.html" % stamp)
    io.open(hp, "w", encoding="utf-8").write(html_doc(SRC))
    for p in (dp, hp):
        print("%s  %.0fKB" % (p, os.path.getsize(p) / 1024))
    kinds = {}
    for k, _v in blocks:
        kinds[k] = kinds.get(k, 0) + 1
    print("  " + " · ".join("%s %d" % (k, n) for k, n in kinds.items()))


main()
