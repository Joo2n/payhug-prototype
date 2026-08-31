#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""용어정의서 워드(.docx) 생성기.

원고는 termsdoc_seed.json 한 곳에서 온다. HTML 편집판(build_termsedit.py)도
같은 원고를 읽으므로 두 산출물은 같은 말을 한다.

산출 — ~/Downloads/payhug_용어정의서/용어정의서_YYYYMMDD.docx

대표가 워드에서 직접 고칠 문서다. 잠금·양식컨트롤을 쓰지 않고 평범한 단락으로만
짠다. 표는 값을 나란히 두는 자리에만 쓴다.
"""

import json
import os
import sys
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PIPE = os.path.dirname(os.path.abspath(__file__))
SEED = os.path.join(PIPE, "termsdoc_seed.json")
OUTDIR = os.path.expanduser("~/Downloads/payhug_용어정의서")

HANGUL = "맑은 고딕"
MONO = "D2Coding"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTE = RGBColor(0x6B, 0x6B, 0x6B)
QUOTE = RGBColor(0x2B, 0x37, 0x4A)
WAIT = RGBColor(0xB4, 0x53, 0x09)
CHECK = RGBColor(0x8A, 0x5A, 0x00)


# ══════════════════════════════════════════════════════════════════
#  글꼴 — 한글은 eastAsia 축을 따로 박아야 적용된다
# ══════════════════════════════════════════════════════════════════

def font(run, name=HANGUL, size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    rf.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def para(doc, text="", size=10.5, bold=False, color=INK, name=HANGUL,
         before=0, after=4, indent=0, line=1.5, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent:
        pf.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    if text:
        font(p.add_run(text), name=name, size=size, bold=bold, color=color)
    return p


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def cell_text(cell, text, size=9.5, bold=False, color=INK, name=HANGUL):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    font(p.add_run(text), name=name, size=size, bold=bold, color=color)


# ══════════════════════════════════════════════════════════════════
#  조각
# ══════════════════════════════════════════════════════════════════

STATUS_COLOR = {"확정": INK, "대기": WAIT, "확인필요": CHECK}


def kv_table(doc, rows):
    """왼쪽 라벨 · 오른쪽 값 2열 표."""
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        r = t.add_row()
        r.cells[0].width = Cm(3.2)
        r.cells[1].width = Cm(12.8)
        shade(r.cells[0], "F4F5F7")
        cell_text(r.cells[0], label, bold=True)
        cell_text(r.cells[1], value)
    para(doc, "", after=6)
    return t


def head(doc, seed):
    para(doc, seed["title"], size=20, bold=True, after=2)
    para(doc, seed["basis"], size=10, color=MUTE, after=14)

    sc = seed.get("scale", {})
    kv_table(doc, [
        ("기준일", seed.get("asof", "")),
        ("규모", "투자실행액 {:,}원 · 순현금 {:,}원 · 가맹점 {}곳 · 할인율 {}%".format(
            sc.get("투자실행액", 0), sc.get("순현금", 0), sc.get("가맹점수", 0),
            ("%g" % (sc.get("할인율", 0) * 100)))),
        ("용도", seed.get("meeting", "")),
    ])


def pending_block(doc, seed):
    items = seed.get("pending") or []
    if not items:
        return
    para(doc, "미결", size=14, bold=True, before=14, after=6)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, (label, w) in enumerate((("항목", 5.4), ("누가", 2.2), ("무엇을", 8.4))):
        hdr.cells[i].width = Cm(w)
        shade(hdr.cells[i], "ECEEF2")
        cell_text(hdr.cells[i], label, bold=True)
    for it in items:
        r = t.add_row()
        cell_text(r.cells[0], it.get("item", ""), bold=True)
        cell_text(r.cells[1], it.get("who", ""))
        cell_text(r.cells[2], it.get("what", ""))
    para(doc, "", after=8)


def pending_formula(seed, it):
    """대기 중인 산식은 원고 한 곳(pending_formula)에서만 온다.

    42항(⑤)은 그 블록을, 45항(⑥)은 그 블록의 depends 를 가리킨다.
    ⑤ 가 오면 블록 한 줄만 고치면 두 항이 함께 바뀐다.
    """
    pf = seed.get("pending_formula")
    if not pf or not it.get("formula_ref"):
        return None
    # 대체 여부는 따로 적어 두지 않고 두 문자열을 견줘 낸다.
    # 그래야 대표 산식이 왔을 때 formula 한 줄만 고쳐도 표시가 맞는다.
    mark = pf.get("status") if pf.get("formula") == pf.get("basis_formula") else None
    if it.get("no") == pf.get("no"):
        out = dict(pf)
    else:
        hit = [d for d in pf.get("depends", []) if d.get("no") == it.get("no")]
        if not hit:
            return None
        out = dict(hit[0])
    out["status"] = mark
    return out


def item_block(doc, seed, it):
    no = it.get("no")
    term = it.get("term", "")
    status = it.get("status", "확정")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    font(p.add_run("%02d  " % no), size=12, bold=True, color=MUTE)
    font(p.add_run(term), size=13.5, bold=True)
    if status != "확정":
        font(p.add_run("   " + status), size=10, bold=True,
             color=STATUS_COLOR.get(status, INK))

    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Cm(0.6)
    q.paragraph_format.space_before = Pt(2)
    q.paragraph_format.space_after = Pt(6)
    q.paragraph_format.line_spacing = 1.4
    font(q.add_run(it.get("quote", "")), size=10, color=QUOTE)

    para(doc, it.get("plain", ""), size=10.5, after=6, line=1.6)

    pf = pending_formula(seed, it)
    rows = []
    if pf:
        rows.append(("산식", pf.get("formula", ""), pf.get("status")))
    elif it.get("formula"):
        rows.append(("산식", it["formula"], None))
    if it.get("symbol"):
        rows.append(("기호", it["symbol"], None))
    if pf and pf.get("our_value"):
        rows.append(("우리 값", pf["our_value"], pf.get("status")))
    elif it.get("our_value"):
        rows.append(("우리 값", it["our_value"], None))
    if it.get("screen"):
        rows.append(("화면", it["screen"], None))
    if it.get("note"):
        rows.append(("비고", it["note"], None))
    if rows:
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for label, value, mark in rows:
            r = t.add_row()
            r.cells[0].width = Cm(2.4)
            r.cells[1].width = Cm(13.6)
            shade(r.cells[0], "F7F8FA")
            cell_text(r.cells[0], label, bold=True, size=9)
            mono = label in ("산식", "기호")
            cell_text(r.cells[1], value, size=9.5,
                      name=(MONO if mono else HANGUL))
            if mark:
                font(r.cells[1].paragraphs[0].add_run("   " + mark),
                     name=HANGUL, size=9, bold=True, color=WAIT)


# ══════════════════════════════════════════════════════════════════
#  조립
# ══════════════════════════════════════════════════════════════════

def build(seed):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for axis in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rf.set(qn(axis), HANGUL)

    head(doc, seed)
    pending_block(doc, seed)

    items = sorted(seed["items"], key=lambda x: x["no"])
    for img in (1, 2):
        group = [i for i in items if i.get("image") == img]
        if not group:
            continue
        doc.add_page_break()
        para(doc, "[%d번 이미지]  %d항" % (img, len(group)),
             size=15, bold=True, before=0, after=4)
        for it in group:
            item_block(doc, seed, it)

    return doc


def main():
    if not os.path.exists(SEED):
        sys.exit("원고 없음 — %s" % SEED)
    with open(SEED, encoding="utf-8") as f:
        seed = json.load(f)

    os.makedirs(OUTDIR, exist_ok=True)
    stamp = os.environ.get("TERMSDOC_STAMP") or date.today().strftime("%Y%m%d")
    out = os.path.join(OUTDIR, "용어정의서_%s.docx" % stamp)
    build(seed).save(out)

    n = len(seed["items"])
    n1 = sum(1 for i in seed["items"] if i.get("image") == 1)
    print("%s\n  항목 %d (1번 %d · 2번 %d) · %.0fKB"
          % (out, n, n1, n - n1, os.path.getsize(out) / 1024))


if __name__ == "__main__":
    main()
