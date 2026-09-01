#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""확정 용어·기호 정리 — 워드와 HTML 두 벌.

원고는 final_terms.json 한 곳에서 온다.
산출 ~/Downloads/payhug_용어정의서/용어기호정리_YYYYMMDD.docx
     _pipeline/investor_admin/final_terms.fragment.html
     payhug-investor-admin/final-terms.html
"""
import html, importlib.util, io, json, os
from datetime import date
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PIPE = os.path.dirname(os.path.abspath(__file__))
SEED = os.path.join(PIPE, "final_terms.json")
REPO = "/Users/semi/cursor/payhug-investor-admin"
OUT = os.path.expanduser("~/Downloads/payhug_용어정의서")
HAN, MONO = "맑은 고딕", "D2Coding"
INK, MUTE, ACC, WARN = RGBColor(0x1A,0x1A,0x1A), RGBColor(0x6B,0x6B,0x6B), RGBColor(0x2C,0x44,0x70), RGBColor(0x9C,0x62,0x12)
e = lambda s: html.escape("" if s is None else str(s))

def font(run, name=HAN, size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    if rf.getparent() is None: rpr.append(rf)
    for ax in ("w:ascii","w:hAnsi","w:eastAsia"): rf.set(qn(ax), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color
    return run

def para(doc, t="", size=10.5, bold=False, color=INK, name=HAN, before=0, after=4, line=1.55):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), line
    if t: font(p.add_run(t), name, size, bold, color)
    return p

def shade(c, hx):
    el = OxmlElement("w:shd"); el.set(qn("w:val"),"clear"); el.set(qn("w:fill"),hx)
    c._tc.get_or_add_tcPr().append(el)

def cell(c, t, size=9.5, bold=False, color=INK, name=HAN):
    c.text = ""; p = c.paragraphs[0]
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    font(p.add_run(t), name, size, bold, color)

def table(doc, head, rows, widths, monos=()):
    t = doc.add_table(rows=1, cols=len(head)); t.style = "Table Grid"
    for i,(h,w) in enumerate(zip(head,widths)):
        t.rows[0].cells[i].width = Cm(w); shade(t.rows[0].cells[i], "ECEEF2")
        cell(t.rows[0].cells[i], h, 9, True, MUTE)
    for r in rows:
        rr = t.add_row()
        for i,v in enumerate(r):
            rr.cells[i].width = Cm(widths[i])
            cell(rr.cells[i], v or "-", 9.5, name=(MONO if i in monos else HAN))
    para(doc, "", after=8)

def docx(d, path):
    doc = Document(); s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.0); s.top_margin = s.bottom_margin = Cm(2.0)
    n = doc.styles["Normal"]; n.font.size = Pt(10.5)
    rpr = n.element.get_or_add_rPr(); rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for ax in ("w:ascii","w:hAnsi","w:eastAsia"): rf.set(qn(ax), HAN)

    para(doc, d["title"], 20, True, after=2)
    para(doc, d["basis"] + " · " + d["asof"], 10, color=MUTE, after=16)

    para(doc, "1. 이름 짓는 규칙", 15, True, before=6, after=6)
    for r in d["rules"]:
        para(doc, "%d. %s" % (r["n"], r["head"]), 11, True, after=2)
        para(doc, r["body"], 10, color=MUTE, after=8)

    para(doc, "2. 범위", 15, True, before=10, after=6)
    table(doc, ["표시","이름","무엇"], [[s["mark"],s["name"],s["def"]] for s in d["scopes"]], [2.2,2.4,12.4], monos=(0,))

    para(doc, "3. 용어와 기호", 15, True, before=10, after=6)
    for kind in ("상수","개념","낱개","집계"):
        vs = [v for v in d["vars"] if v["kind"] == kind]
        if not vs: continue
        para(doc, kind, 11.5, True, color=ACC, before=6, after=4)
        for v in vs:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
            font(p.add_run(v["term"] + "  "), size=11.5, bold=True)
            font(p.add_run(v["sym"]), MONO, 11, True, ACC)
            if v.get("formula"):
                para(doc, v["formula"], 9.5, name=MONO, color=ACC, after=2)
            para(doc, v["plain"], 10, color=MUTE, after=4)

    para(doc, "4. 계산 예시", 15, True, before=12, after=4)
    para(doc, d["calc"]["기준"], 10, color=MUTE, after=6)
    table(doc, ["단계","값"], [[a,b] for a,b in d["calc"]["steps"]], [10.6,6.4], monos=(1,))
    para(doc, "금액으로 되돌린 검산", 11.5, True, before=6, after=4)
    table(doc, ["단계","값"], [[a,b] for a,b in d["calc"]["검산"]], [10.6,6.4], monos=(1,))

    para(doc, "5. 확인 대기", 15, True, before=12, after=6)
    table(doc, ["항목","무엇이 걸리나"], [[p["item"],p["note"]] for p in d["pending"]], [4.6,12.4])
    doc.save(path)

CSS = """
:root{--bg:#FAF9F7;--surface:#FFF;--sunk:#F4F2ED;--ink:#1B1A17;--mute:#6B6862;
 --rule:#E5E1D9;--rule-hard:#CFC9BD;--accent:#2C4470;--accent-soft:#2C447014;
 --warn:#9C6212;--warn-soft:#9C621214;--shadow:0 1px 2px #1b1a1710,0 8px 22px #1b1a170a}
@media (prefers-color-scheme:dark){:root:not([data-theme="light"]){
 --bg:#16161A;--surface:#1D1E23;--sunk:#22242A;--ink:#ECEAE5;--mute:#9B978E;
 --rule:#2E3038;--rule-hard:#3D4048;--accent:#93AEE0;--accent-soft:#93AEE01F;
 --warn:#E0A64A;--warn-soft:#E0A64A20;--shadow:0 1px 2px #00000040,0 8px 22px #00000030}}
:root[data-theme="dark"]{--bg:#16161A;--surface:#1D1E23;--sunk:#22242A;--ink:#ECEAE5;
 --mute:#9B978E;--rule:#2E3038;--rule-hard:#3D4048;--accent:#93AEE0;--accent-soft:#93AEE01F;
 --warn:#E0A64A;--warn-soft:#E0A64A20;--shadow:0 1px 2px #00000040,0 8px 22px #00000030}
*{box-sizing:border-box}
body{margin:0;background:var(--bg);color:var(--ink);font-size:15px;line-height:1.78;
 font-family:"Noto Sans KR",-apple-system,BlinkMacSystemFont,"Apple SD Gothic Neo",sans-serif;
 -webkit-font-smoothing:antialiased}
code,.m{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:.92em}
.wrap{max-width:900px;margin:0 auto;padding:34px 20px 110px}
h1{margin:0 0 5px;font-family:"Noto Serif KR",serif;font-size:30px;font-weight:700;letter-spacing:-.02em}
.meta{color:var(--mute);font-size:13px;margin:0 0 26px}
h2{margin:38px 0 10px;padding-top:22px;border-top:2px solid var(--rule-hard);
 font-family:"Noto Serif KR",serif;font-size:20px;font-weight:700}
h3{margin:22px 0 8px;font-size:14px;font-weight:700;color:var(--accent);letter-spacing:.04em}
p{margin:0 0 12px}
.rule{background:var(--surface);border:1px solid var(--rule);border-radius:10px;
 padding:13px 16px;margin:9px 0;box-shadow:var(--shadow)}
.rule b{display:block;font-size:15px;margin-bottom:4px}
.rule span{color:var(--mute);font-size:13.5px;line-height:1.7}
.scroll{overflow-x:auto;border:1px solid var(--rule);border-radius:9px;margin:12px 0}
table{width:100%;border-collapse:collapse;font-size:13.5px}
th,td{padding:9px 12px;text-align:left;border-bottom:1px solid var(--rule);vertical-align:top}
th{background:var(--sunk);font-size:11.5px;letter-spacing:.04em;color:var(--mute);font-weight:700;white-space:nowrap}
tr:last-child td{border-bottom:0}
td.n{font-family:"IBM Plex Mono",ui-monospace,monospace;white-space:nowrap}
.v{background:var(--surface);border:1px solid var(--rule);border-radius:10px;
 padding:12px 15px;margin:9px 0;box-shadow:var(--shadow)}
.v .hd{display:flex;align-items:baseline;gap:9px;flex-wrap:wrap}
.v .t{font-weight:700;font-size:15px}
.v .s{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:12.5px;
 color:var(--accent);background:var(--accent-soft);padding:2px 8px;border-radius:5px}
.v .f{margin:9px 0 0;padding:9px 12px;background:var(--sunk);border-radius:7px;
 font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:12.5px;line-height:1.7;
 white-space:pre-wrap;overflow-x:auto}
.v .p{margin:8px 0 0;font-size:13.5px;line-height:1.72;color:var(--mute)}
.warn{background:var(--warn-soft);border:1px solid var(--warn);border-radius:10px;padding:13px 16px;margin:12px 0}
.warn b{color:var(--warn)}
@media print{.rule,.v{break-inside:avoid;box-shadow:none}}
"""

def html_doc(d):
    o = ['<title>%s</title>' % e(d["title"]),
      '<style>@import url("https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;500;700&family=Noto+Serif+KR:wght@400;700&family=IBM+Plex+Mono:wght@400;500&display=swap");',
      CSS, "</style>", '<div class="wrap">']
    o.append('<h1>%s</h1><p class="meta">%s · %s</p>' % (e(d["title"]), e(d["basis"]), e(d["asof"])))

    o.append("<h2>1. 이름 짓는 규칙</h2>")
    for r in d["rules"]:
        o.append('<div class="rule"><b>%d. %s</b><span>%s</span></div>'
                 % (r["n"], e(r["head"]), e(r["body"]).replace("`","")))

    o.append("<h2>2. 범위</h2><div class='scroll'><table><thead><tr><th>표시</th><th>이름</th><th>무엇</th></tr></thead><tbody>")
    for s in d["scopes"]:
        o.append("<tr><td class='n'>%s</td><td>%s</td><td>%s</td></tr>" % (e(s["mark"]), e(s["name"]), e(s["def"]).replace("`","")))
    o.append("</tbody></table></div>")

    o.append("<h2>3. 용어와 기호</h2>")
    for kind in ("상수","개념","낱개","집계"):
        vs = [v for v in d["vars"] if v["kind"] == kind]
        if not vs: continue
        o.append("<h3>%s</h3>" % kind)
        for v in vs:
            o.append('<div class="v"><div class="hd"><span class="t">%s</span><span class="s">%s</span></div>'
                     % (e(v["term"]), e(v["sym"])))
            if v.get("formula"): o.append('<div class="f">%s</div>' % e(v["formula"]))
            o.append('<p class="p">%s</p></div>' % e(v["plain"]).replace("`",""))

    o.append("<h2>4. 계산 예시</h2><p class='meta'>%s</p>" % e(d["calc"]["기준"]))
    o.append("<div class='scroll'><table><thead><tr><th>단계</th><th>값</th></tr></thead><tbody>")
    for a,b in d["calc"]["steps"]:
        o.append("<tr><td>%s</td><td class='n'>%s</td></tr>" % (e(a), e(b)))
    o.append("</tbody></table></div><h3>금액으로 되돌린 검산</h3>")
    o.append("<div class='scroll'><table><thead><tr><th>단계</th><th>값</th></tr></thead><tbody>")
    for a,b in d["calc"]["검산"]:
        o.append("<tr><td>%s</td><td class='n'>%s</td></tr>" % (e(a), e(b)))
    o.append("</tbody></table></div>")

    o.append("<h2>5. 확인 대기</h2>")
    o.append("<div class='scroll'><table><thead><tr><th>항목</th><th>무엇이 걸리나</th></tr></thead><tbody>")
    for p in d["pending"]:
        o.append("<tr><td>%s</td><td>%s</td></tr>" % (e(p["item"]), e(p["note"]).replace("`","")))
    o.append("</tbody></table></div></div>")
    return "".join(o)

def main():
    d = json.load(io.open(SEED, encoding="utf-8"))
    os.makedirs(OUT, exist_ok=True)
    stamp = os.environ.get("TERMSDOC_STAMP") or date.today().strftime("%Y%m%d")
    dp = os.path.join(OUT, "용어기호정리_%s.docx" % stamp)
    docx(d, dp)
    frag = html_doc(d)
    spec = importlib.util.spec_from_file_location("bt", os.path.join(PIPE,"build_termsedit.py"))
    bt = importlib.util.module_from_spec(spec); spec.loader.exec_module(bt)
    fp = os.path.join(PIPE, "final_terms.fragment.html")
    hp = os.path.join(REPO, "final-terms.html")
    io.open(fp,"w",encoding="utf-8").write(frag)
    io.open(hp,"w",encoding="utf-8").write(bt.full(frag))
    for p in (dp, fp, hp): print("%s  %.0fKB" % (p, os.path.getsize(p)/1024))
    print("  규칙 %d · 범위 %d · 용어 %d · 확인대기 %d"
          % (len(d["rules"]), len(d["scopes"]), len(d["vars"]), len(d["pending"])))

main()
