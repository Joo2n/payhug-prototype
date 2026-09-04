#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""워드에서 고친 용어·기호 정리를 원고 꼴로 되읽는다.

대상   ~/Downloads/payhug_용어정의서/용어기호정리_*.docx  (build_final.py 가 낸 것)
견줌   final_terms.json

원고에 쓰지 않는다. 차이 목록만 낸다.

  python3 read_wordedit.py                     최신 워드를 찾아 차이만 낸다
  python3 read_wordedit.py <파일.docx>         파일을 지정
  python3 read_wordedit.py --out <경로.json>   되읽은 원고를 딴 파일로 뽑는다
  python3 read_wordedit.py --all               고친 것이 없어도 전건을 훑는다

워드는 표와 문단이라 어느 글이 어느 칸인지 서식으로 가른다.
  5절 용어 이름·기호   굵은 11.5pt 로 시작하는 문단. 두 칸 띄어쓰기 뒤가 기호
  5절 산식             9.5pt D2Coding 문단
  5절 설명             10pt 흐린 글자 문단
  1절 기존 표기        첫 표의 첫 칸
  3절 규칙             11pt 굵은 문단과 그 아래 10pt 문단
  4절 범위             세 번째 표

되읽기의 한계 — build_final.py 가 설명을 워드에 넣을 때 원고의 백틱을 뗀다.
그래서 워드에서 되읽은 설명에는 백틱이 없다. 견줄 때 원고 쪽 백틱도 떼고 본다.
"""

import io
import json
import os
import re
import sys

from docx import Document

PIPE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PIPE)
import read_finaledit as rf

SEED_PATH = os.path.join(PIPE, "final_terms.json")
OUTDIR = os.path.expanduser("~/Downloads/payhug_용어정의서")
DOWNLOADS = os.path.expanduser("~/Downloads")
STEM = "용어기호정리"

MONO = "D2Coding"
SEC5 = "5. 용어와 기호"
SEC3 = "3. 이름 짓는 규칙"
SEC4 = "4. 범위"
SEC6 = "6. 계산 예시"


# ══════════════════════════════════════════════════════════════════
#  워드에서 글자를 되살린다
# ══════════════════════════════════════════════════════════════════

def runtext(p):
    """문단·칸 하나를 원고 표기로 되돌린다. 아래첨자는 `_i` · `_{d−1}` 로 돌아온다."""
    out = []
    for r in p.runs:
        t = r.text
        if r.font.subscript:
            # 워드 조판은 쉼표 뒤에 가는 공백(U+2009)을 넣는다. 원고에는 없다.
            b = t.replace(",\u2009", ",").replace(", ", ",")
            out.append("_" + b if b in ("i", "r") else "_{" + b + "}")
        else:
            out.append(t)
    return "".join(out)


def celltext(c):
    return "\n".join(runtext(p) for p in c.paragraphs).strip("\n")


def pt(run):
    return run.font.size.pt if run.font.size else None


def kind(p):
    """문단이 무엇인지 서식으로 가른다."""
    rs = [r for r in p.runs if r.text]
    if not rs:
        return None
    head, s = rs[0], pt(rs[0])
    if s == 15:
        return "sec"
    if s == 11.5 and head.bold:
        return "term"
    if s == 11 and head.bold and head.font.name != MONO:
        return "rulehead"
    if s == 9.5 and any(r.font.name == MONO for r in rs):
        return "formula"
    if s == 10:
        return "plain"
    return None


def read_docx(path):
    """워드 한 벌을 final_terms.json 과 같은 꼴로 되돌린다."""
    doc = Document(path)
    ps = doc.paragraphs

    out = {"title": "", "basis": "", "asof": "", "rules": [], "scopes": [], "vars": []}
    if ps:
        out["title"] = runtext(ps[0]).strip()
    if len(ps) > 1:
        meta = runtext(ps[1]).strip()
        # 머리글은 `근거 · 기준일` 한 줄이다. 마지막 가운뎃점 뒤가 기준일이다.
        m = re.match(r"^(.*) · (\d{4}-\d{2}-\d{2})$", meta)
        out["basis"], out["asof"] = (m.group(1), m.group(2)) if m else (meta, "")

    at = {}
    for i, p in enumerate(ps):
        t = p.text.strip()
        for key, head in (("s3", SEC3), ("s4", SEC4), ("s5", SEC5), ("s6", SEC6)):
            if t.startswith(head):
                at.setdefault(key, i)

    # ── 3절 이름 짓는 규칙 ────────────────────────────────
    if "s3" in at and "s4" in at:
        cur = None
        for p in ps[at["s3"] + 1:at["s4"]]:
            k = kind(p)
            if k == "rulehead":
                txt = runtext(p).strip()
                m = re.match(r"^(\d+)\.\s*(.*)$", txt)
                cur = {"n": int(m.group(1)) if m else len(out["rules"]) + 1,
                       "head": m.group(2) if m else txt, "body": ""}
                out["rules"].append(cur)
            elif k == "plain" and cur is not None:
                cur["body"] = (cur["body"] + " " + runtext(p).strip()).strip()

    # ── 4절 범위 · 1절 기존 표기 — 표에서 읽는다 ───────────
    tbs = doc.tables
    alias_rows = []
    if len(tbs) >= 1:
        for r in tbs[0].rows[1:]:
            cs = [celltext(c) for c in r.cells]
            if len(cs) >= 3 and (cs[0] or cs[1]):
                alias_rows.append((cs[0], cs[1], cs[2]))
    if len(tbs) >= 3:
        for r in tbs[2].rows[1:]:
            cs = [celltext(c) for c in r.cells]
            if len(cs) >= 3 and cs[0]:
                out["scopes"].append({"mark": cs[0], "name": cs[1], "def": cs[2]})

    # ── 5절 용어와 기호 ───────────────────────────────────
    if "s5" in at:
        stop = at.get("s6", len(ps))
        cur = None
        for p in ps[at["s5"] + 1:stop]:
            k = kind(p)
            if k == "term":
                txt = runtext(p)
                m = re.split(r"\s{2,}", txt, 1)
                term = m[0].strip()
                sym = m[1].strip() if len(m) > 1 else ""
                if not sym:
                    # 두 칸 띄어쓰기가 지워졌으면 D2Coding 런에서 가른다
                    a, b = [], []
                    hit = False
                    for r in p.runs:
                        if r.font.name == MONO:
                            hit = True
                        (b if hit else a).append(r)
                    term = "".join(r.text for r in a).strip()
                    sym = "".join(r.text for r in b).strip()
                cur = {"term": term, "sym": sym, "alias": None,
                       "formula": None, "plain": ""}
                out["vars"].append(cur)
            elif cur is None:
                continue
            elif k == "formula":
                cur["formula"] = runtext(p)
            elif k == "plain":
                cur["plain"] = (cur["plain"] + " " + runtext(p).strip()).strip()

    # ── 기존 표기를 항에 붙인다 — 기호로, 안 맞으면 용어 이름으로 ──
    by_sym = {a[1]: a[0] for a in alias_rows if a[1]}
    by_term = {a[2]: a[0] for a in alias_rows if a[2]}
    for v in out["vars"]:
        v["alias"] = by_sym.get(v["sym"]) or by_term.get(v["term"]) or None

    return out


def find_latest():
    cand = []
    for d in (OUTDIR, DOWNLOADS):
        if not os.path.isdir(d):
            continue
        for f in os.listdir(d):
            if re.match(r"^%s_\d{8}" % STEM, f) and f.endswith(".docx"):
                p = os.path.join(d, f)
                cand.append((os.path.getmtime(p), p))
    return max(cand)[1] if cand else None


# ══════════════════════════════════════════════════════════════════

def bare(s):
    """워드는 백틱을 떼고 찍는다. 견줄 때 원고 쪽도 떼고 본다."""
    return "" if s is None else str(s).replace("`", "")


def main():
    args = list(sys.argv[1:])
    out_path = None
    if "--out" in args:
        i = args.index("--out")
        out_path = args[i + 1]
        del args[i:i + 2]
    show_all = "--all" in args
    args = [a for a in args if not a.startswith("--")]

    src = args[0] if args else find_latest()
    if not src or not os.path.exists(src):
        sys.exit("워드를 못 찾음 — %s 아래 %s_*.docx" % (OUTDIR, STEM))

    new = read_docx(src)
    old = json.load(io.open(SEED_PATH, encoding="utf-8"))

    print("워드    %s" % src)
    print("원고    %s" % SEED_PATH)
    print("용어    원고 %d항 → 워드 %d항"
          % (len(old.get("vars", [])), len(new.get("vars", []))))
    print("규칙 %d · 범위 %d · 기존 표기 %d줄"
          % (len(new["rules"]), len(new["scopes"]),
             sum(1 for v in new["vars"] if v.get("alias"))))
    print()

    facts = rf.fact_strings()
    banned = rf.banned_words()

    rows = rf.diff(old, new, prep=bare)
    rows += rf.diff_side(old, new, "rules", "규칙", ("head", "body"), prep=bare)
    rows += rf.diff_side(old, new, "scopes", "범위", ("name", "def"), prep=bare)
    n = rf.show(rows, facts, banned, "차이")

    risk = rf.verifier_risk(rows)
    if risk:
        print()
        print("verify_final_terms.py 가 볼 자리 %d곳" % len(risk))
        for no, term, field, a, b in risk:
            print("  [%02d] %s  %s" % (no, term, field))
        print("  검증기 기준을 이 값으로 옮길지 정해 주면 그때 고친다.")

    if show_all:
        print()
        print("전건 훑기 — 금지어·원장에 없는 숫자")
        hit = 0
        for i, v in enumerate(new.get("vars", []), 1):
            for f in rf.VAR_FIELDS:
                t = rf.norm(v.get(f))
                h = [w for w in banned if w in t]
                nums = rf.scan_numbers(t, facts)
                if h or nums:
                    hit += 1
                    print("  [%02d] %s  %s  %s" % (i, rf.norm(v.get("term")), f,
                                                   " · ".join(h + nums)))
        if not hit:
            print("  없음")

    if out_path:
        io.open(out_path, "w", encoding="utf-8").write(
            json.dumps(new, ensure_ascii=False, indent=1) + "\n")
        print()
        print("되읽은 원고  %s" % out_path)

    return n


if __name__ == "__main__":
    main()
